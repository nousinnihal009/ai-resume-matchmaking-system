"""
Text extraction pipeline for processing resume and job documents.
"""
import io
from typing import Optional
import structlog

try:
    from PyPDF2 import PdfReader
    import docx
except ImportError:
    PdfReader = None
    docx = None

logger = structlog.get_logger(__name__)


async def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text from uploaded file (PDF or DOCX).

    Args:
        file_content: Raw file bytes
        filename: Original filename

    Returns:
        Extracted text content

    Raises:
        ValueError: If file type is unsupported or extraction fails
    """
    if not file_content:
        raise ValueError("Empty file content")

    file_extension = filename.lower().split('.')[-1]

    logger.info("Extracting text from file", filename=filename, extension=file_extension)

    try:
        if file_extension == 'pdf':
            return _extract_text_from_pdf(file_content)
        elif file_extension in ['docx', 'doc']:
            return _extract_text_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    except Exception as e:
        logger.error("Text extraction failed", error=str(e), filename=filename)
        raise ValueError(f"Failed to extract text: {str(e)}")


def _extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file.
    """
    if PdfReader is None:
        raise ImportError("PyPDF2 is required for PDF processing")

    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)

        text_content = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text.strip():
                text_content.append(text)

        full_text = '\n'.join(text_content)

        if not full_text.strip():
            raise ValueError("No readable text found in PDF")

        logger.info("PDF text extracted", pages=len(pdf_reader.pages), text_length=len(full_text))
        return full_text

    except Exception as e:
        logger.error("PDF extraction failed", error=str(e))
        raise ValueError(f"PDF extraction failed: {str(e)}")


def _extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file.
    """
    if docx is None:
        raise ImportError("python-docx is required for DOCX processing")

    try:
        docx_file = io.BytesIO(file_content)
        doc = docx.Document(docx_file)

        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text)

        full_text = '\n'.join(text_content)

        if not full_text.strip():
            raise ValueError("No readable text found in DOCX")

        logger.info("DOCX text extracted", text_length=len(full_text))
        return full_text

    except Exception as e:
        logger.error("DOCX extraction failed", error=str(e))
        raise ValueError(f"DOCX extraction failed: {str(e)}")


async def extract_text_from_job_description(description: str) -> str:
    """
    Extract and clean text from job description.
    """
    if not description:
        return ""

    # Basic cleaning - remove extra whitespace
    cleaned = ' '.join(description.split())
    logger.info("Job description text extracted", text_length=len(cleaned))
    return cleaned


def extract_basic_info_from_text(text: str) -> dict:
    """
    Extract basic information from text (name, email, phone, etc.).
    This is a simple pattern-based extraction.
    """
    import re

    info = {
        "name": None,
        "email": None,
        "phone": None,
        "education": [],
        "experience": [],
    }

    # Extract email
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    if email_match:
        info["email"] = email_match.group()

    # Extract phone (simple pattern)
    phone_pattern = r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'
    phone_match = re.search(phone_pattern, text)
    if phone_match:
        info["phone"] = phone_match.group()

    # Extract education keywords
    education_keywords = [
        "bachelor", "master", "phd", "doctorate", "associate", "diploma",
        "university", "college", "institute", "school"
    ]

    lines = text.split('\n')
    for line in lines:
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in education_keywords):
            info["education"].append(line.strip())

    # Extract experience (simple year-based detection)
    experience_pattern = r'\b(19|20)\d{2}\b'
    years = set(re.findall(experience_pattern, text))
    if years:
        info["experience_years"] = sorted(list(years))

    return info
